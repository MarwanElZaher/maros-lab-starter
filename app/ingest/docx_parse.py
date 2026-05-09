"""
docx → paragraphs (plain text) + embedded image OCR texts.

Returns two things:
  - paragraphs: list[str]  — all paragraph texts (for vector store)
  - image_texts: list[dict]  — OCR results from embedded images
"""
from __future__ import annotations
from pathlib import Path


def parse_docx(docx_path: Path) -> tuple[list[str], list[dict]]:
    """
    Returns (paragraphs, image_texts).
    paragraphs: list of non-empty paragraph strings.
    image_texts: list of {para_idx, text} from OCR (may be empty if tesseract absent).
    """
    from docx import Document
    from .image_ocr import extract_image_texts

    doc = Document(str(docx_path))
    paragraphs: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also grab table cell text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text and text not in paragraphs:
                    paragraphs.append(text)

    image_texts = extract_image_texts(docx_path)
    return paragraphs, image_texts
